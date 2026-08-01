"""Markdown -> DOCX conversion for generated documents.

Handles the subset of markdown our drafting prompts emit: #/##/### headings,
paragraphs, bullet lists, bold/italic runs, and [^n] footnote markers (kept
inline as superscript-style text). Times New Roman 12pt, the convention for
USCIS filings.
"""
import io
import re

from docx import Document
from docx.shared import Pt


def _add_runs(paragraph, text: str) -> None:
    # Split on **bold** and *italic* while keeping the markers' content.
    for token in re.split(r"(\*\*.+?\*\*|\*.+?\*)", text):
        if not token:
            continue
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("*") and token.endswith("*"):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        else:
            paragraph.add_run(token)


def markdown_to_docx(md: str, title: str | None = None) -> bytes:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    if title:
        doc.add_heading(title, level=0)

    footnotes: list[str] = []
    lines = md.split("\n")
    for raw in lines:
        line = raw.rstrip()
        if not line.strip():
            continue
        m = re.match(r"^(#{1,4})\s+(.*)$", line)
        if m:
            doc.add_heading(re.sub(r"\*\*?", "", m.group(2)), level=len(m.group(1)))
            continue
        m = re.match(r"^\s*[-*]\s+(.*)$", line)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            _add_runs(p, m.group(1))
            continue
        m = re.match(r"^\s*\d+\.\s+(.*)$", line)
        if m:
            p = doc.add_paragraph(style="List Number")
            _add_runs(p, m.group(1))
            continue
        m = re.match(r"^\[\^(\d+)\]:\s*(.*)$", line)
        if m:
            footnotes.append(f"[{m.group(1)}] {m.group(2)}")
            continue
        p = doc.add_paragraph()
        _add_runs(p, line)

    if footnotes:
        doc.add_paragraph("_" * 30)
        for note in footnotes:
            p = doc.add_paragraph()
            p.style = doc.styles["Normal"]
            run = p.add_run(note)
            run.font.size = Pt(10)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
